"""Generación de informe de auditoría pentest en formato Word (.docx)."""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.core.config import REPORT_AUDITOR_NAME, SCREENSHOTS_DIR
from app.services.report_content import (
    build_audit_steps,
    executive_summary,
    ports_summary,
    risk_level,
    safe_target,
)
from app.services.step_capture import capture_all_evidence


def _find_screenshot(target: str, step_num: int, step_key: str) -> Optional[Path]:
    base = SCREENSHOTS_DIR / safe_target(target)
    if not base.is_dir():
        return None

    exts = (".png", ".jpg", ".jpeg", ".webp")
    patterns = [f"{step_num:02d}_{step_key}", f"{step_num:02d}", step_key]
    if step_key == "flag":
        patterns.extend(["99_flag", "flag"])

    for pattern in patterns:
        for ext in exts:
            path = base / f"{pattern}{ext}"
            if path.is_file():
                return path
    return None


def _add_evidence_block(
    doc: Document,
    target: str,
    step_num: int,
    step_key: str,
    caption: str,
) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Evidencia gráfica (captura automática ReconTool)\n").bold = True

    image_path = _find_screenshot(target, step_num, step_key)
    if image_path:
        doc.add_picture(str(image_path), width=Inches(5.5))
        cap = doc.add_paragraph(f"Figura {step_num}: {caption}")
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
        return

    placeholder = doc.add_paragraph("[Captura en generación — reintente descargar el informe]")
    placeholder.runs[0].italic = True
    placeholder.runs[0].font.color.rgb = RGBColor(0x99, 0x66, 0x00)


def _add_label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}\n")
    run_label.bold = True
    run_label.font.size = Pt(11)
    p.add_run(value).font.size = Pt(11)


def _add_step_block(
    doc: Document,
    step_num: int,
    title: str,
    tool: str,
    objective: str,
    findings: str,
    target: str,
    step_key: str,
) -> None:
    doc.add_heading(f"Paso {step_num}  —  {title}", level=2)
    _add_label_value(doc, "Herramienta", tool)
    _add_label_value(doc, "Objetivo", objective)
    _add_label_value(doc, "Hallazgos / Resultado", findings)
    _add_evidence_block(doc, target, step_num, step_key, title)
    doc.add_paragraph()


def generate_audit_report(scan_data: dict, auditor: str | None = None) -> bytes:
    if not scan_data.get("evidence_screenshots"):
        capture_all_evidence(scan_data)

    target = scan_data.get("target", "objetivo")
    auditor = auditor or REPORT_AUDITOR_NAME
    now = datetime.now()
    date_str = now.strftime("%d / %m / %Y")
    risk = risk_level(scan_data)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p_conf = doc.add_paragraph("CONFIDENCIAL  ·  USO RESTRINGIDO")
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_conf.runs[0].font.size = Pt(9)
    p_conf.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    t1 = doc.add_heading("INFORME TÉCNICO", level=0)
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2 = doc.add_heading("AUDITORÍA DE SEGURIDAD OFENSIVA", level=1)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _add_label_value(doc, "Objetivo", target)
    _add_label_value(doc, "Puertos analizados", ports_summary(scan_data))
    _add_label_value(doc, "Fecha de auditoría", date_str)
    _add_label_value(doc, "Nivel de riesgo", risk)
    _add_label_value(doc, "Clasificación", "CONFIDENCIAL — LABORATORIO AUTORIZADO")
    _add_label_value(doc, "Auditor", auditor)

    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "Generado automáticamente por ReconTool con evidencias capturadas en cada fase del pipeline."
    )
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].italic = True

    footer = doc.add_paragraph(f"ReconTool Pentesting Report · v1.0 · {now.year}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)

    doc.add_page_break()

    doc.add_heading("ÍNDICE", level=1)
    for item in [
        "1. Resumen Ejecutivo",
        "2. Metodología y Alcance",
        "3. Desarrollo de la Auditoría",
        "4. Resumen de Vulnerabilidades",
        "5. Recomendaciones",
        "6. Resultado — Flag obtenida",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    doc.add_heading("1. Resumen Ejecutivo", level=1)
    doc.add_paragraph(executive_summary(scan_data))

    doc.add_heading("2. Metodología y Alcance", level=1)
    doc.add_paragraph(
        "Metodología basada en PTES y OWASP: reconocimiento, enumeración, "
        "correlación CVE y explotación controlada en lab."
    )

    doc.add_heading("3. Desarrollo de la Auditoría", level=1)
    doc.add_paragraph("Cada paso incluye captura automática generada por ReconTool.")

    audit_steps = build_audit_steps(scan_data)
    for idx, step in enumerate(audit_steps, start=1):
        _add_step_block(
            doc,
            idx,
            step["title"],
            step["tool"],
            step["objective"],
            step["findings"],
            target,
            step.get("key", f"step{idx}"),
        )

    pipeline = scan_data.get("pipeline_log") or []
    if pipeline:
        doc.add_heading("Cronología del pipeline", level=2)
        for entry in pipeline:
            msg = entry.get("message", "")
            line = f"{entry.get('phase')}: {entry.get('status')}"
            if msg:
                line += f" — {msg}"
            doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("4. Resumen de Vulnerabilidades", level=1)
    vuln_table = doc.add_table(rows=1, cols=4)
    vuln_table.style = "Table Grid"
    for i, h in enumerate(["ID", "VULNERABILIDAD", "SEVERIDAD", "CVSS"]):
        vuln_table.rows[0].cells[i].text = h

    vid = 1
    seen: set[str] = set()
    for row in scan_data.get("results", []):
        for f in row.get("cve_findings") or []:
            if "error" in f:
                continue
            cve = f.get("cve", "")
            if not cve or cve in seen:
                continue
            seen.add(cve)
            r = vuln_table.add_row().cells
            r[0].text = f"V-{vid:02d}"
            r[1].text = f"{f.get('description', cve)[:120]}"
            r[2].text = str(f.get("severity", "—"))
            r[3].text = str(f.get("cvss", "—"))
            vid += 1

    if scan_data.get("exploitation", {}).get("flag_captured"):
        r = vuln_table.add_row().cells
        r[0].text = f"V-{vid:02d}"
        r[1].text = "FTP anónimo — lectura de flag.txt"
        r[2].text = "CRÍTICO"
        r[3].text = "9.1"

    doc.add_heading("5. Recomendaciones", level=1)
    for rec in [
        "Corregir misconfiguración FTP / login anónimo.",
        "No exponer archivos sensibles en servicios de transferencia.",
        "Parchear servicios con CVEs detectados.",
    ]:
        doc.add_paragraph(rec, style="List Bullet")

    doc.add_heading("6. Resultado — Flag obtenida", level=1)
    flags = scan_data.get("exploitation", {}).get("flags") or []
    if flags:
        for f in flags:
            p = doc.add_paragraph()
            p.add_run(f"Archivo: {f.get('filename')}\n").bold = True
            run_flag = p.add_run(f.get("content", ""))
            run_flag.font.name = "Consolas"
            run_flag.font.size = Pt(12)
            run_flag.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
        _add_evidence_block(doc, target, 99, "flag", "Flag capturada — evidencia final")
    else:
        doc.add_paragraph("No se capturó flag en esta ejecución.")

    doc.add_paragraph()
    end = doc.add_paragraph("— FIN DEL INFORME —")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.runs[0].bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
