#!/usr/bin/env python3
"""Genera el informe Word de la Práctica 1 — Ciberseguridad con IA."""

from __future__ import annotations

import io
import textwrap
from datetime import date
from pathlib import Path

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "informe-practica1" / "assets"
OUTPUT = ROOT / "docs" / "informe-practica1" / "Informe_Practica1_ReconTool.docx"
RECON_LOGO = ROOT / "code" / "frontend" / "img" / "recontool-logo.png"
EVOLVE_LOGO = ASSETS / "evolve-logo.png"

AUTHOR = "[Nombre y apellidos del alumno]"
GITHUB = "https://github.com/ionut5251/Proyecto-ReconTool.git"
FONT_PATHS = [
    "C:/Windows/Fonts/consola.ttf",
    "C:/Windows/Fonts/segoeui.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
]


def _font(size: int, mono: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    paths = ["C:/Windows/Fonts/consola.ttf"] if mono else FONT_PATHS
    for path in paths:
        if Path(path).is_file():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def _download_evolve_logo() -> Path | None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    candidates = [
        "https://evolve.es/wp-content/uploads/2024/06/logo-evolve.svg",
        "https://evolve.es/wp-content/uploads/2023/03/logo-evolve.svg",
        "https://evolve.es/wp-content/themes/evolve/assets/img/logo.svg",
    ]
    try:
        with httpx.Client(timeout=20.0, follow_redirects=True) as client:
            html = client.get("https://evolve.es/").text
            import re

            for url in re.findall(r'https://evolve\.es[^"\']+\.(?:png|svg|webp)', html):
                candidates.insert(0, url)
    except Exception:
        pass

    for url in candidates:
        try:
            with httpx.Client(timeout=20.0, follow_redirects=True) as client:
                resp = client.get(url)
                if resp.status_code != 200 or len(resp.content) < 200:
                    continue
                ext = ".png" if "png" in url else ".svg"
                path = ASSETS / f"evolve-logo{ext}"
                path.write_bytes(resp.content)
                if ext == ".svg":
                    return None
                return path
        except Exception:
            continue

    # Placeholder tipográfico si no hay logo descargable
    img = Image.new("RGB", (480, 120), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    font = _font(52)
    draw.text((40, 30), "EVOLVE", fill=(20, 20, 20), font=font)
    draw.text((40, 85), "Academy · Master Ciberseguridad e IA", fill=(100, 100, 100), font=_font(14))
    img.save(EVOLVE_LOGO)
    return EVOLVE_LOGO


def _render_ui_mock(title: str, lines: list[str], filename: str) -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    width, height = 960, 540
    img = Image.new("RGB", (width, height), (10, 10, 10))
    draw = ImageDraw.Draw(img)
    font = _font(16, mono=True)
    font_title = _font(22)

    draw.rectangle((0, 0, width, 56), fill=(26, 92, 56))
    draw.text((20, 14), title, fill=(255, 255, 255), font=font_title)

    y = 80
    for line in lines:
        color = (180, 255, 180) if "HTB{" in line or "flag" in line.lower() else (220, 220, 220)
        draw.text((24, y), line[:110], fill=color, font=font)
        y += 26

    out = ASSETS / filename
    img.save(out)
    return out


def _render_architecture_diagram() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    w, h = 1000, 620
    img = Image.new("RGB", (w, h), (248, 249, 250))
    draw = ImageDraw.Draw(img)
    font = _font(15)
    font_b = _font(17)

    def box(x, y, bw, bh, label, fill=(255, 255, 255), border=(61, 143, 98)):
        draw.rounded_rectangle((x, y, x + bw, y + bh), radius=8, fill=fill, outline=border, width=2)
        draw.text((x + 12, y + bh // 2 - 10), label, fill=(30, 30, 30), font=font_b)

    draw.text((320, 20), "Arquitectura ReconTool", fill=(30, 30, 30), font=_font(24))
    box(380, 70, 240, 60, "Frontend (HTML/JS)")
    box(380, 200, 240, 60, "FastAPI /api/*")
    box(80, 330, 200, 55, "nmap")
    box(310, 330, 200, 55, "CVE (NVD + local)")
    box(540, 330, 200, 55, "OSINT pasivo")
    box(770, 330, 180, 55, "IA operativa")
    box(180, 450, 180, 55, "FTP / Telnet / SMB")
    box(420, 450, 180, 55, "Informes Word/HTML")
    box(660, 450, 180, 55, "Capturas evidencia")

    def arrow(x1, y1, x2, y2):
        draw.line((x1, y1, x2, y2), fill=(61, 143, 98), width=2)

    arrow(500, 130, 500, 200)
    arrow(500, 260, 500, 310)
    arrow(430, 310, 180, 330)
    arrow(500, 310, 410, 330)
    arrow(570, 310, 640, 330)
    arrow(650, 310, 860, 330)
    arrow(180, 385, 270, 450)
    arrow(500, 385, 510, 450)
    arrow(860, 385, 750, 450)

    out = ASSETS / "diagrama-arquitectura.png"
    img.save(out)
    return out


def _render_process_diagram() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    w, h = 1000, 280
    img = Image.new("RGB", (w, h), (248, 249, 250))
    draw = ImageDraw.Draw(img)
    font = _font(14)
    font_b = _font(16)
    steps = [
        "Fase 1\nRecon pasivo",
        "Plan de\nataque + IA",
        "Fase 2\nAtaque activo",
        "Flag +\nInforme",
    ]
    x = 40
    for i, label in enumerate(steps):
        draw.rounded_rectangle((x, 90, x + 200, 190), radius=10, fill=(255, 255, 255), outline=(61, 143, 98), width=2)
        for j, line in enumerate(label.split("\n")):
            draw.text((x + 20, 115 + j * 22), line, fill=(30, 30, 30), font=font_b if j == 0 else font)
        if i < len(steps) - 1:
            draw.line((x + 200, 140, x + 240, 140), fill=(61, 143, 98), width=3)
            draw.polygon([(x + 240, 140), (x + 228, 132), (x + 228, 148)], fill=(61, 143, 98))
        x += 240
    draw.text((40, 30), "Pipeline operativo — dos fases", fill=(30, 30, 30), font=_font(22))
    out = ASSETS / "diagrama-proceso.png"
    img.save(out)
    return out


def _add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld)
    doc.add_page_break()


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _body(doc: Document, text: str) -> None:
    for para in text.strip().split("\n\n"):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(8)


def _bullet(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_image(doc: Document, path: Path, width: float = 5.8, caption: str = "") -> None:
    if path.is_file():
        doc.add_picture(str(path), width=Inches(width))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        doc.add_paragraph(f"[Imagen pendiente: {path.name}]")


def build_document() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    evolve = _download_evolve_logo()
    arch = _render_architecture_diagram()
    process = _render_process_diagram()
    ui_home = _render_ui_mock(
        "ReconTool — Búsqueda",
        [
            "Reconocimiento de vulnerabilidades.",
            "",
            "¿Qué IP quieres analizar?",
            "[ 10.129.x.x                    ] [ Escanear ]",
            "",
            "Master Ciberseguridad",
        ],
        "ui-home.png",
    )
    ui_passive = _render_ui_mock(
        "ReconTool — Fase 1 pasiva",
        [
            "Target: 10.129.x.x",
            "Recon pasivo completado — 1 puerto(s), N CVE(s)",
            "",
            "Puertos: 445/tcp SMB | OS: Linux",
            "Vector detectado: SMB — shares anónimos",
            "",
            "[ Proceder con ataque activo ]",
        ],
        "ui-fase1.png",
    )
    ui_active = _render_ui_mock(
        "ReconTool — Fase 2 activa",
        [
            "Ataque activo: smb_anonymous (ai/heuristic)",
            "[list_shares] smbclient -L //IP -N",
            "[explore_share] Share: ...",
            "[mget] 12 archivo(s) descargados",
            "",
            "FLAG capturada: HTB{...}",
            "[ Informe Word ] [ Informe HTML ]",
        ],
        "ui-fase2.png",
    )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- PORTADA ---
    for _ in range(2):
        doc.add_paragraph()
    if evolve and evolve.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(evolve), width=Inches(2.8))
    else:
        t = doc.add_paragraph("EVOLVE")
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.runs[0].font.size = Pt(36)
        t.runs[0].bold = True

    doc.add_paragraph()
    if RECON_LOGO.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(RECON_LOGO), width=Inches(1.4))

    doc.add_paragraph()
    title = doc.add_paragraph("ReconTool")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].bold = True

    sub = doc.add_paragraph(
        "Herramienta de reconocimiento y pentesting asistido por IA\n"
        "Línea Red Team · Módulo Ciberseguridad Avanzada"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)

    doc.add_paragraph()
    meta = doc.add_paragraph(
        f"Práctica 1 — Desarrollo de Herramienta de Ciberseguridad con IA\n"
        f"Máster en Ciberseguridad e Inteligencia Artificial · Evolve\n"
        f"Autor: {AUTHOR}\n"
        f"Fecha: {date.today().strftime('%d / %m / %Y')}\n"
        f"Repositorio: {GITHUB}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(11)
    meta.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # --- ÍNDICE ---
    _heading(doc, "Índice de contenidos", 1)
    toc_items = [
        "1. Portada",
        "2. Índice de contenidos",
        "3. Resumen ejecutivo",
        "4. Descripción del problema y justificación de la solución",
        "5. Arquitectura técnica",
        "6. Proceso de desarrollo (fase a fase)",
        "7. Guía de despliegue",
        "8. Manual de uso de la herramienta",
        "9. Conclusiones y lecciones aprendidas",
        "10. Road map de mejora (Práctica 2)",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # --- 3 RESUMEN EJECUTIVO ---
    _heading(doc, "3. Resumen ejecutivo", 1)
    _body(
        doc,
        """ReconTool es una herramienta web de reconocimiento ofensivo y apoyo a pentesting desarrollada en el marco del Máster en Ciberseguridad e Inteligencia Artificial. El proyecto combina escaneo de servicios (nmap), correlación de vulnerabilidades (base local y API NVD), OSINT pasivo y sugerencia de vectores de ataque mediante modelos de lenguaje (OpenAI u Ollama local), con un módulo de explotación controlada en entornos autorizados.

La herramienta responde a la necesidad de automatizar las fases iniciales de una auditoría ofensiva — reconocimiento, priorización de hallazgos y documentación — sin sustituir el criterio del analista. El flujo se organiza en dos fases: recon pasivo (puertos, CVEs, OSINT y plan de ataque) y ataque activo bajo confirmación del usuario, ejecutando solo el vector detectado (FTP anónimo, Telnet, SMB, etc.).

El desarrollo se ha realizado con asistencia de IA (Cursor) siguiendo arquitectura modular en Python/FastAPI y frontend estático. Toda la documentación viva del proyecto reside en la bóveda Obsidian bovrecon/, interconectada con guías de API, pipeline, despliegue en Kali y roadmap. El código está versionado en GitHub.

En la versión actual, desplegada localmente (127.0.0.1:8000) por limitaciones económicas del servidor Hetzner, la herramienta genera informes de auditoría en Word e HTML con capturas automáticas de evidencia. Se han validado vectores sobre laboratorios controlados (FTP, Telnet y SMB). La evolución hacia despliegue cloud y agente IA plenamente autónomo constituye el siguiente hito del proyecto.""",
    )
    doc.add_page_break()

    # --- 4 PROBLEMA ---
    _heading(doc, "4. Descripción del problema y justificación de la solución", 1)
    _body(
        doc,
        """En auditorías de seguridad y laboratorios de Red Team, el reconocimiento inicial consume una parte significativa del tiempo: identificar puertos abiertos, asociar CVEs, decidir el siguiente paso y documentar hallazgos. Las herramientas existentes (nmap, searchsploit, smbclient, etc.) son potentes pero fragmentadas; el analista debe correlacionar manualmente resultados y redactar informes.

ReconTool unifica este flujo en un panel web accesible, inspirado en interfaces tipo Shodan pero orientado a pentesting ético en scope acordado. La integración de IA permite priorizar vectores cuando coexisten varios servicios (por ejemplo FTP y SMB) y generar recomendaciones en lenguaje natural, mientras un motor operativo (ai_operational.py) elige el módulo de explotación más adecuado.

Se eligió la línea Red Team (práctica #07 del enunciado: reconocimiento activo automatizado) por alineación directa con el temario del máster y la posibilidad de demostrar un pipeline completo: recon → CVE → vector → flag → informe. El uso exclusivo en entornos autorizados (VPN de laboratorio, VMs propias) garantiza el cumplimiento ético y legal.""",
    )

    # --- 5 ARQUITECTURA ---
    _heading(doc, "5. Arquitectura técnica", 1)
    _body(
        doc,
        """La solución sigue una arquitectura en capas: frontend estático (HTML/CSS/JS), API REST FastAPI, servicios de dominio modulares y dependencias externas (nmap, NVD, LLM, smbclient en Kali).

Componentes principales:
• Frontend: index.html (búsqueda) y results.html (resultados en dos fases).
• API: /api/scan (pasivo), /api/attack (activo), /api/report (informes).
• Pipeline: scan_pipeline.py orquesta nmap, CVE, OSINT, plan de ataque y explotación.
• Módulos ofensivos: ftp_probe, telnet_probe, smb_probe ejecutados vía exploit_runner.
• IA: ai_advisor (vectores) y ai_operational (elección de módulo).
• Documentación: bóveda Obsidian bovrecon/ con más de 25 notas enlazadas.""",
    )
    _add_image(doc, arch, caption="Figura 1. Diagrama de arquitectura lógica de ReconTool.")
    _body(
        doc,
        "El diagrama de flujo de datos (secuencia usuario → API → nmap → CVE → IA) está documentado en bovrecon/02-Arquitectura/Flujo-de-datos.md y puede exportarse desde Obsidian como figura adicional si se desea sustituir la imagen anterior.",
    )

    # --- 6 PROCESO ---
    _heading(doc, "6. Proceso de desarrollo (fase a fase)", 1)
    _add_image(doc, process, caption="Figura 2. Pipeline operativo en dos fases.")
    phases = [
        (
            "Fase A — Diseño y arquitectura",
            "Definición de visión Red Team, estructura code/backend + code/frontend, bóveda bovrecon y repositorio GitHub.",
        ),
        (
            "Fase B — Recon pasivo",
            "Integración nmap, enriquecimiento CVE (NVD + JSON local), OSINT pasivo (reverse DNS, probe HTTP) y UI de resultados.",
        ),
        (
            "Fase C — IA y vectores",
            "Servicio ai_advisor con OpenAI/Ollama, playbook de respaldo y panel de vectores sugeridos en la UI.",
        ),
        (
            "Fase D — Explotación modular",
            "Módulos FTP, Telnet y SMB; router de servicios; fase activa desacoplada con botón «Proceder con ataque».",
        ),
        (
            "Fase E — IA operativa",
            "ai_operational.py elige módulo cuando hay varios puertos abiertos; heurística telnet → smb → ftp → http.",
        ),
        (
            "Fase F — Informes y evidencias",
            "Generación Word/HTML con logo, capturas automáticas por paso (Pillow) y descarga desde la UI.",
        ),
        (
            "Fase G — Documentación y pruebas",
            "Pruebas en Kali con VPN de laboratorio; script recontool para arranque; commits incrementales en GitHub.",
        ),
    ]
    for title, desc in phases:
        _heading(doc, title, 2)
        _body(doc, desc)

    _heading(doc, "Evidencias de interfaz (capturas representativas)", 2)
    _body(doc, "A continuación se incluyen mockups generados del estado actual de la UI. Sustituir por capturas reales del navegador si se desea mayor fidelidad visual.")
    _add_image(doc, ui_home, width=5.2, caption="Figura 3. Pantalla de búsqueda.")
    _add_image(doc, ui_passive, width=5.2, caption="Figura 4. Fase 1 — recon pasivo y plan de ataque.")
    _add_image(doc, ui_active, width=5.2, caption="Figura 5. Fase 2 — ataque activo e informes.")

    _heading(doc, "Herramientas de IA utilizadas en el desarrollo", 2)
    _bullet(
        doc,
        [
            "Cursor — IDE con agente IA para implementación, refactor y documentación.",
            "Modelos LLM — OpenAI GPT-4o-mini y/o Ollama (llama3.2) en runtime de la herramienta.",
            "Obsidian — documentación viva interconectada (bovrecon/).",
        ],
    )

    # --- 7 DESPLIEGUE ---
    _heading(doc, "7. Guía de despliegue paso a paso", 1)
    _body(
        doc,
        "Nota: el enunciado recomienda Hetzner Cloud. Por restricciones económicas actuales el despliegue documentado es local/Kali; la migración a VPS es directa (mismos pasos sin VPN de lab).",
    )
    _heading(doc, "7.1 Requisitos", 2)
    _bullet(
        doc,
        [
            "Python 3.10+, nmap en PATH, smbclient (Kali: samba-common-bin).",
            "Privilegios elevados para nmap -sS -O (root o Administrator).",
            "Opcional: OPENAI_API_KEY o Ollama local para IA.",
        ],
    )
    _heading(doc, "7.2 Instalación (Windows — desarrollo)", 2)
    for line in [
        'cd "Proyecto ReconTool\\code\\backend"',
        "python -m venv venv",
        ".\\venv\\Scripts\\Activate.ps1",
        "pip install -r requirements.txt",
        "copy .env.example .env",
        "python run.py",
        "Abrir http://127.0.0.1:8000",
    ]:
        doc.add_paragraph(line, style="List Number")

    _heading(doc, "7.3 Instalación (Kali — pruebas ofensivas)", 2)
    for line in [
        "git clone https://github.com/ionut5251/Proyecto-ReconTool.git",
        "cd Proyecto-ReconTool && git pull",
        "chmod +x recontool && bash recontool",
        "Conectar VPN del laboratorio antes de escanear IPs internas.",
    ]:
        doc.add_paragraph(line, style="List Number")

    _heading(doc, "7.4 Despliegue futuro en Hetzner (planificado)", 2)
    _bullet(
        doc,
        [
            "Provisionar VPS Ubuntu, instalar nmap, Python, smbclient.",
            "Clonar repo, configurar .env y systemd o Docker.",
            "Exponer puerto 8000 con reverse proxy (nginx) y HTTPS.",
            "Dominio propio apuntando a la IP del VPS.",
        ],
    )

    # --- 8 MANUAL ---
    _heading(doc, "8. Manual de uso de la herramienta", 1)
    _body(
        doc,
        """Esta sección describe el uso básico de ReconTool. El alumno puede ampliarla con capturas reales del flujo completo.

1. Acceder a la URL del servidor (local: http://127.0.0.1:8000).
2. Introducir la IP o hostname autorizado y pulsar «Escanear».
3. Revisar la Fase 1: puertos, CVEs, OSINT, vector detectado y análisis IA/playbook.
4. Pulsar «Proceder con ataque activo» para ejecutar el módulo correspondiente (FTP, Telnet, SMB…).
5. Si se captura flag, descargar informe Word (.docx) o HTML desde el panel final.
6. Consultar bovrecon/ para detalle técnico de API, módulos y configuración .env.

[Espaacio reservado para capturas adicionales del manual — completar por el autor.]""",
    )

    # --- 9 CONCLUSIONES ---
    _heading(doc, "9. Conclusiones y lecciones aprendidas", 1)
    _body(
        doc,
        """ReconTool demuestra que es viable construir, con apoyo de IA en el desarrollo y en runtime, una herramienta de Red Team modular y documentada profesionalmente en el plazo de una práctica académica.

Lecciones aprendidas:
• La separación recon pasivo / ataque activo mejora la claridad operativa y reduce ejecuciones no deseadas.
• La correlación CVE automática aporta valor, pero requiere filtrado para evitar ruido en servicios genéricos.
• La IA es útil para priorizar vectores, pero debe complementarse con módulos deterministas y heurísticas de respaldo.
• La documentación continua (Obsidian + Git) facilita retomar el proyecto tras pausas y preparar la entrega.

Limitaciones actuales: despliegue local en lugar de Hetzner, módulos de explotación limitados a servicios concretos, y dependencia de herramientas del sistema (nmap, smbclient) en el host donde corre el backend.

Con tiempo y esfuerzo, el proyecto puede evolucionar hacia un agente autónomo multi-fase (recon → explotación → post-explotación), despliegue cloud accesible desde Internet y cumplimiento íntegro de todos los requisitos formales de la práctica. La base técnica, el repositorio y la documentación ya soportan esa extensión.""",
    )

    # --- 10 ROADMAP placeholder ---
    _heading(doc, "10. Road map de mejora (Práctica 2)", 1)
    _body(
        doc,
        "[Sección reservada — completar por el autor con el road map visual exigido en el requisito 05 del enunciado: nuevas funcionalidades, rendimiento, seguridad, integraciones y estimaciones.]",
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Informe generado: {path}")
